# daily_ai_automation.py - Enhanced with advanced sources (Clean version)
import feedparser
import json
import requests
from datetime import datetime  # Rimosso timedelta non usato
import os
import logging
from typing import List, Dict, Any
# Rimosso time e re non usati
from urllib.parse import urlparse
import hashlib
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import textwrap
import uuid
import base64
import openai

load_dotenv()

def generate_post_with_ai(article_data):
    """Genera post LinkedIn usando AI"""
    try:
        logging.info(f"Generazione post per: {article_data.get('title', 'Articolo senza titolo')}")
        
        # Prepara il prompt per l'AI
        prompt = f"""
Crea un post LinkedIn professionale e coinvolgente basato su questo articolo AI:

Titolo: {article_data.get('title', '')}
Riassunto: {article_data.get('summary', '')}
Link: {article_data.get('link', '')}

# Regole per la Creazione di Post LinkedIn Autentici sull'IA

Lunghezza post: 500-600 parole, 10 hashtag

## Principi Fondamentali

### 1. Trasparenza Assoluta
- **Mai inventare esperienze professionali** o progetti non realmente vissuti
- **Mai presentare scenari ipotetici come casi reali** - utilizzare sempre linguaggio che chiarisca il carattere esplorativo
- **Ammettere quando si riparte da fonti esterne** (articoli, studi, conversazioni) senza fingere esperienza diretta
- **Riconoscere limiti e incertezze** delle proprie valutazioni
- **Evitare affermazioni definitive** su implementazioni non verificate

### 2. Anti-Artificialità Totale
- **Zero formule preconfezionate** - ogni post deve sembrare una riflessione spontanea
- **Bandire linguaggio da IA generativa**: emoji sistematiche, punti elenco rigidi, CTA formulaiche
- **Eliminare frasi da "thought leader"**: "Come professionista...", "La mia esperienza mi ha insegnato...", "Ho osservato che..."
- **Priorità alla conversazione naturale** rispetto all'ottimizzazione per engagement
- **Variabilità strutturale** - nessuno schema fisso da seguire

3. Connessione con la Vita Reale
- **Partire sempre da situazioni quotidiane** che i lettori possano riconoscere e vivere
- **Usare esempi concreti dalla vita comune**: cercare sintomi online, usare app, ricevere pubblicità personalizzate
- **Toccare paure e speranze reali** delle persone normali, non solo professionisti del settore
- **Evitare astrazioni** come "framework etici" o "innovazione tecnologica" senza collegamento pratico
- **Focalizzarsi su impatti tangibili** nella vita di tutti i giorni

## Stile e Formato

### Linguaggio
- **Conversazionale e diretto** - come parlare con un amico al bar, non con un collega in riunione
- **Eliminare gergo professionale** o spiegarlo attraverso esempi quotidiani
- **Ritmo naturale** - alternare frasi brevi e lunghe per fluidità di lettura
- **Linguaggio universale** - comprensibile a chiunque, non solo agli addetti ai lavori


### Struttura Flessibile
- **Sviluppo organico** - seguire il flusso naturale del pensiero
- **Lunghezza determinata dal contenuto** - non da regole artificiali
- **Conclusioni aperte** - domande genuine che nascono dalla riflessione
- **Connessione emotiva** - toccare sentimenti che i lettori provano realmente

### Elementi da Evitare Sempre
- ❌ "Come professionista...", "La mia esperienza...", "Ho osservato che..."
- ❌ Progetti inventati o collaborazioni non esistenti
- ❌ Liste di benefici standardizzate
- ❌ Emoji come elementi strutturali
- ❌ Hashtag generici e scontati (#AI #Innovation #Future)
- ❌ Call-to-action commerciali evidenti
- ❌ Tone of voice da "thought leader"
- ❌ Astrazioni senza collegamento alla vita reale
- ❌ Promesse irrealistiche o hype tecnologico

## Obiettivo Finale

### Risultati Desiderati
Creare contenuti che:
- **Facciano sentire i lettori compresi** nelle loro preoccupazioni quotidiane
- **Rendano accessibili concetti complessi** attraverso esempi dalla vita comune
- **Generino riflessione autentica** su come la tecnologia impatta la vita di tutti
- **Stimolino conversazioni significative** basate su esperienze condivise
- **Dimostrino competenza attraverso chiarezza** nella spiegazione, non autoreferenzialità
- **Mantengano credibilità** attraverso onestà e trasparenza

### Impatto sulla Community
- **Democratizzare la comprensione dell'IA** rendendola accessibile a non esperti
- **Creare consapevolezza** su questioni che toccano tutti nella vita quotidiana  
- **Ispirare riflessioni pratiche** su come proteggersi e usare meglio la tecnologia
- **Contribuire a decisioni informate** nella vita personale e professionale
- **Ridurre il divario** tra esperti e persone comuni sui temi tecnologici

## Approccio Provocatorio (Intellettuale) - Versione Autentica

### Filosofia di Base
**Trasformare l'approccio provocatorio da "guru che fa domande retoriche" a "persona normale che condivide autenticamente le proprie scoperte e preoccupazioni sulla tecnologia che ci circonda".**

### Caratteristiche Distintive

#### 1. Vulnerabilità Umana (Non Professionale)
- **Ammettere di non aver capito qualcosa** fino a poco tempo fa
- **Condividere momenti di consapevolezza** genuini
- **Riconoscere di essere stati ingenui** su alcuni aspetti tecnologici
- **Esporre preoccupazioni personali** concrete, non astratte

#### 2. Conflitti Interiori Quotidiani
- **Comodità vs Privacy**: "Mi piace che Spotify mi suggerisca musica perfetta, ma mi inquieta sapere quanto sa di me"
- **Efficienza vs Umanità**: "È comodo il chatbot bancario, ma mi manca parlare con una persona vera"
- **Progresso vs Controllo**: "Questi strumenti mi semplificano la vita, ma mi sento sempre più dipendente"

#### 3. Scoperte che Cambiano Prospettiva
- **Momenti di realizzazione**: "Non sapevo che ogni volta che..."
- **Connessioni inaspettate**: "Mi sono accorto che..."
- **Cambi di opinione**: "Prima pensavo che fosse normale, ora..."

#### 4. Domande che Nascono dalla Vita Reale
- **Preoccupazioni concrete**: "Ma se mio figlio usa questo chatbot per i compiti..."
- **Dilemmi quotidiani**: "Vale la pena rinunciare alla privacy per questa comodità?"
- **Conseguenze pratiche**: "Cosa succede se sbaglia nella diagnosi?"

### Struttura Tipica dell'Approccio Provocatorio Autentico

#### Apertura dalla Vita Quotidiana (30-50 parole)
- Situazione concreta vissuta personalmente o osservata
- Momento di consapevolezza o sorpresa
- Domanda nata da un'esperienza reale
- Contraddizione notata nel comportamento quotidiano

#### Sviluppo attraverso Esempi Concreti (100-150 parole)
- Spiegazione del meccanismo tecnologico in termini semplici
- Esempi che tutti possono riconoscere dalla propria vita
- Implicazioni pratiche immediate
- Connessioni con preoccupazioni comuni

#### Provocazione Autentica (50-80 parole)
- Domande che nascono da paure o speranze reali
- Paradossi che emergono dall'uso quotidiano della tecnologia
- Dilemmi che tutti affrontiamo senza rendercene conto
- Conseguenze che potrebbero toccare la vita di chiunque

#### Chiusura Inclusiva (30-40 parole)
- Domanda che invita alla condivisione di esperienze simili
- Riconoscimento che altri potrebbero avere vissuto situazioni analoghe
- Apertura al confronto su soluzioni pratiche
- Invito a riflettere insieme su scelte quotidiane


### Note Operative Fondamentali
- **Mai inventare esperienze** - basarsi solo su situazioni realmente vissute o verosimilmente osservabili
- **Sempre partire dal quotidiano** - non da articoli tecnici o trend di settore
- **Spiegare in termini semplici** - come se parlassi a un familiare non esperto
- **Toccare emozioni reali** - preoccupazioni, speranze, frustrazioni che tutti provano
- **Evitare pose da esperto** - mantenere il punto di vista della persona comune che riflette
"""
        
        # Simulazione generazione AI (da sostituire con vera API)
        ai_response = generate_ai_content(prompt)
        
        if ai_response:
            logging.info("Post generato con successo")
            return ai_response
        else:
            logging.error("Errore nella generazione del post")
            return None
            
    except Exception as e:
        logging.error(f"Errore nella generazione post: {e}")
        return None

def generate_ai_content(prompt):
    """Genera contenuto usando API AI (Claude/OpenAI)"""
    try:
        # Prova prima con Claude
        claude_key = os.getenv('CLAUDE_API_KEY')
        if claude_key:
            return generate_with_claude(prompt, claude_key)
        
        # Fallback su OpenAI
        openai_key = os.getenv('OPENAI_API_KEY')
        if openai_key:
            return generate_with_openai(prompt, openai_key)
        
        # Fallback su contenuto template
        logging.warning("Nessuna API key trovata, usando template")
        return generate_template_content(prompt)
        
    except Exception as e:
        logging.error(f"Errore generazione AI: {e}")
        return generate_template_content(prompt)

def generate_with_claude(prompt, api_key):
    """Genera contenuto con Claude"""
    try:
        import anthropic
        
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        return response.content[0].text
        
    except Exception as e:
        logging.error(f"Errore Claude API: {e}")
        return None

def generate_with_openai(prompt, api_key):
    """Genera contenuto con OpenAI"""
    try:
        import openai
        
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        logging.error(f"Errore OpenAI API: {e}")
        return None

def generate_template_content(prompt):
    """Genera contenuto template quando non ci sono API keys"""
    # Estrai titolo dal prompt
    title = "Novità nell'Intelligenza Artificiale"
    if "Titolo:" in prompt:
        title = prompt.split("Titolo:")[1].split("\n")[0].strip()
    
    template = f"""🤖 {title}

L'intelligenza artificiale continua a evolversi rapidamente, portando innovazioni che trasformano il nostro modo di lavorare e vivere.

🎯 Applicazioni pratiche:
• Automatizzazione di processi ripetitivi
• Analisi di grandi volumi di dati
• Miglioramento dell'esperienza utente
• Supporto alle decisioni aziendali

💡 Questa tecnologia offre opportunità straordinarie per aumentare la produttività e creare nuove soluzioni innovative.

Cosa ne pensi? Come vedi l'AI integrata nel tuo settore?

#AI #IntelligenzaArtificiale #Innovazione #Tecnologia #Futuro"""
    
    return template

def create_post_image(article_data, post_content, provider="openai"):
    """Crea immagine AI per il post usando OpenAI o Google"""
    try:
        logging.info(f"Generazione immagine AI per il post usando {provider.upper()}")
        
        # Genera prompt per l'immagine basato sull'articolo
        title = article_data.get('title', 'AI Innovation')
        summary = article_data.get('summary', 'Latest AI technology advancement')
        
        image_prompt = f"""Create a professional, modern image for a LinkedIn post about: {title}. 
The image should be:
- Professional and business-appropriate for LinkedIn
- Related to artificial intelligence and technology
- Clean, modern design with tech/AI elements
- Suitable for social media sharing
- High quality and visually appealing

Context: {summary[:200]}

Style: Professional, tech-focused, modern graphics, suitable for LinkedIn business post."""
        
        # Nome file temporaneo
        temp_filename = f"temp_ai_image_{uuid.uuid4().hex[:8]}.png"
        
        success = False
        
        if provider.lower() == "google":
            # Usa Google Gemini
            google_key = os.getenv('GOOGLE_API_KEY')
            if not google_key:
                logging.warning("API key Google non trovata, provo con OpenAI")
                provider = "openai"
            else:
                success = generate_post_image_google(image_prompt, temp_filename, google_key)
        
        if provider.lower() == "openai" and not success:
            # Usa OpenAI (fallback o scelta primaria)
            openai_key = os.getenv('OPENAI_API_KEY')
            if not openai_key:
                logging.warning("API key OpenAI non trovata, uso immagine template")
                return create_template_image(article_data, post_content)
            success = generate_post_image_ai(image_prompt, temp_filename, openai_key)
        
        if success and os.path.exists(temp_filename):
            logging.info(f"Immagine AI generata con {provider.upper()}: {temp_filename}")
            return temp_filename
        else:
            logging.warning("Generazione AI fallita, uso immagine template")
            return create_template_image(article_data, post_content)
            
    except Exception as e:
        logging.error(f"Errore generazione immagine AI: {e}")
        return create_template_image(article_data, post_content)

def generate_post_image_google(image_generation_prompt, output_filepath, api_key):
    """
    Genera un'immagine basata su un prompt testuale usando Google Gemini API
    """
    logging.info(f"Richiesta generazione immagine Google AI: '{image_generation_prompt[:100]}...'")
    
    try:
        from google import genai
        from google.genai import types
        from PIL import Image
        from io import BytesIO
        
        # Configura il client Google Gemini
        genai.configure(api_key=api_key)
        client = genai.Client()
        
        # Genera l'immagine usando Gemini
        response = client.models.generate_content(
            model="gemini-2.5-flash-image-preview",
            contents=[image_generation_prompt],
        )
        
        # Gestione della risposta - estrai l'immagine generata
        if response and hasattr(response, 'candidates') and response.candidates:
            candidate = response.candidates[0]
            if hasattr(candidate, 'content') and candidate.content.parts:
                for part in candidate.content.parts:
                    if hasattr(part, 'inline_data') and part.inline_data:
                        # Decodifica i dati dell'immagine
                        image_data = part.inline_data.data
                        
                        # Se i dati sono in base64, decodificali
                        if isinstance(image_data, str):
                            import base64
                            image_bytes = base64.b64decode(image_data)
                        else:
                            image_bytes = image_data
                        
                        # Salva l'immagine
                        with open(output_filepath, 'wb') as f:
                            f.write(image_bytes)
                        
                        logging.info(f"Immagine Google AI salvata: {output_filepath}")
                        return True
        
        logging.error("Risposta Google API non contiene immagine generata")
        if response:
            logging.error(f"Struttura risposta: {type(response)}")
        return False
        
    except ImportError as e:
        logging.error(f"Librerie Google Gemini non installate: {e}. Installare con: pip install google-generativeai")
        return False
    except Exception as e:
        logging.error(f"Errore durante generazione immagine Google AI: {e}")
        return False

def generate_post_image_ai(image_generation_prompt, output_filepath, api_key):
    """
    Genera un'immagine basata su un prompt testuale usando OpenAI DALL-E 3 API
    """
    logging.info(f"Richiesta generazione immagine AI: '{image_generation_prompt[:100]}...'")

    try:
        client = openai.OpenAI(api_key=api_key)

        # Usa l'API corretta di OpenAI per la generazione di immagini con DALL-E 3
        response = client.images.generate(
            model="dall-e-3",
            prompt=image_generation_prompt,
            size="1024x1024",
            quality="standard",
            n=1
        )

        # Estrae l'URL dell'immagine generata
        image_url = response.data[0].url

        # Scarica l'immagine dall'URL
        import requests
        image_response = requests.get(image_url, timeout=30)
        image_response.raise_for_status()

        # Salva l'immagine nel file
        with open(output_filepath, 'wb') as f:
            f.write(image_response.content)

        logging.info(f"Immagine AI salvata: {output_filepath}")

        # Log del prompt rivisto se disponibile
        if hasattr(response.data[0], 'revised_prompt') and response.data[0].revised_prompt:
            logging.info(f"Prompt migliorato dall'AI: '{response.data[0].revised_prompt[:100]}...'")

        return True

    except Exception as e:
        logging.error(f"Errore API OpenAI per generazione immagine: {e}")
        return False

def create_template_image(article_data, post_content):
    """Crea immagine template quando API AI non è disponibile (fallback)"""
    try:
        logging.info("Creazione immagine template (fallback)")
        
        # Dimensioni ottimali per LinkedIn
        width, height = 1200, 630
        
        # Colori moderni per AI/Tech
        bg_color = '#0a0e27'  # Blu scuro tech
        accent_color = '#00d4ff'  # Ciano brillante
        text_color = '#ffffff'
        
        # Crea immagine
        image = Image.new('RGB', (width, height), color=bg_color)
        draw = ImageDraw.Draw(image)
        
        # Font (prova font migliori)
        try:
            title_font = ImageFont.truetype("arial.ttf", 42)
            subtitle_font = ImageFont.truetype("arial.ttf", 24)
            brand_font = ImageFont.truetype("arial.ttf", 18)
        except Exception:
            title_font = ImageFont.load_default()
            subtitle_font = ImageFont.load_default()
            brand_font = ImageFont.load_default()
        
        # Titolo dell'articolo (max 2 righe)
        title = article_data.get('title', 'AI Innovation News')
        title_lines = textwrap.wrap(title, width=35)
        
        # Disegna titolo
        y_pos = 120
        for line in title_lines[:2]:
            draw.text((60, y_pos), line, fill=text_color, font=title_font)
            y_pos += 50
        
        # Linea decorativa
        draw.rectangle([60, y_pos + 20, 400, y_pos + 25], fill=accent_color)
        
        # Categoria/Fonte
        source = article_data.get('source', 'AI News').upper()
        draw.text((60, y_pos + 50), f"🤖 {source}", fill=accent_color, font=subtitle_font)
        
        # Brand/footer
        draw.text((60, height - 80), "AI Content Helper • LinkedIn Post", fill='#888888', font=brand_font)
        
        # Elemento grafico decorativo (cerchi tech)
        for i in range(3):
            x = width - 200 + i * 30
            y = 150 + i * 40
            draw.ellipse([x, y, x + 15, y + 15], fill=accent_color)
        
        # Salva immagine template
        temp_path = f"temp_template_{uuid.uuid4().hex[:8]}.png"
        image.save(temp_path)
        
        logging.info(f"Immagine template creata: {temp_path}")
        return temp_path
        
    except Exception as e:
        logging.error(f"Errore creazione immagine template: {e}")
        return None

def save_post_to_docx(post_content, article_data, image_path=None):
    """Salva post in formato Word"""
    try:
        # Crea cartella per data
        today = datetime.now()
        date_folder = f"generated_posts/{today.year}/{today.month:02d}/{today.day:02d}"
        os.makedirs(date_folder, exist_ok=True)
        
        # Crea documento Word
        doc = Document()
        
        
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Data: {today.strftime('%Y-%m-%d %H:%M')}\n").bold = True
        meta_para.add_run(f"Fonte: {article_data.get('source', 'N/A')}\n").bold = True
        meta_para.add_run(f"Link: {article_data.get('link', 'N/A')}\n").bold = True
        
        # Contenuto post
        doc.add_heading('Contenuto Post LinkedIn:', 2)
        doc.add_paragraph(post_content)
        
        # Immagine (se presente)
        if image_path and os.path.exists(image_path):
            doc.add_heading('Immagine:', 2)
            doc.add_picture(image_path, width=Inches(6))
        
        # Salva documento
        timestamp = today.strftime('%H%M%S')
        filename = f"post_{timestamp}.docx"
        filepath = os.path.join(date_folder, filename)
        
        doc.save(filepath)
        logging.info(f"Post salvato: {filepath}")
        
        #Salva immagine permanente invece di cancellarla
        if image_path and os.path.exists(image_path):
            import shutil
            image_filename = f"post_{timestamp}.png"
            permanent_image_path = os.path.join(date_folder, image_filename)
            shutil.copy2(image_path, permanent_image_path)
            os.remove(image_path)  # Rimuovi immagine temporanea
            logging.info(f"Immagine salvata: {permanent_image_path}")
        
        return filepath
        
    except Exception as e:
        logging.error(f"Errore salvataggio post: {e}")
        return None

def main():
    """Funzione principale dell'automazione"""
    try:
        logging.info("=== Avvio Automazione AI Enhanced ===")
        
        # Inizializza collector enhanced
        collector = EnhancedArticleCollector()
        
        # Raccoglie articoli da tutte le fonti
        articles = collector.collect_all_articles()
        
        if not articles:
            logging.warning("Nessun articolo raccolto")
            return
        
        # Seleziona i migliori articoli
        selected_articles = collector.select_top_articles(articles)
        
        # Genera post per ogni articolo selezionato
        generated_posts = []
        for i, article in enumerate(selected_articles, 1):
            logging.info(f"Generazione post {i}/{len(selected_articles)}")
            
            # Genera contenuto post
            post_content = generate_post_with_ai(article)
            
            if post_content:
                # Crea immagine
                image_path = create_post_image(article, post_content)
                
                # Salva post
                saved_path = save_post_to_docx(post_content, article, image_path)
                
                if saved_path:
                    generated_posts.append(saved_path)
                    logging.info(f"Post {i} completato: {saved_path}")
                else:
                    logging.error(f"Errore salvataggio post {i}")
            else:
                logging.error(f"Errore generazione contenuto post {i}")
        
        # Riepilogo finale
        logging.info(f"Automazione completata. Generati {len(generated_posts)} post")
        
        if generated_posts:
            logging.info("Post generati:")
            for post_path in generated_posts:
                logging.info(f"  - {post_path}")
        
        return generated_posts
        
    except Exception as e:
        logging.error(f"Errore nell'automazione: {e}")
        return []

class EnhancedArticleCollector:
    def __init__(self, config_path: str = "automation_config.json"):
        self.config = self.load_config(config_path)
        self.setup_logging()
        self.processed_articles = set()
        
    def load_config(self, config_path: str) -> Dict[str, Any]:
        """Carica configurazione da file JSON"""
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            logging.error(f"File di configurazione {config_path} non trovato")
            return {}
    
    def setup_logging(self):
        """Setup logging system"""
        logs_dir = "logs"
        if not os.path.exists(logs_dir):
            os.makedirs(logs_dir)
        
        log_file = os.path.join(logs_dir, f"enhanced_automation_{datetime.now().strftime('%Y%m%d')}.log")
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file, encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
    
    def get_article_hash(self, article: Dict[str, Any]) -> str:
        """Genera hash unico per l'articolo per evitare duplicati"""
        content = f"{article.get('title', '')}{article.get('link', '')}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def fetch_rss_articles(self) -> List[Dict[str, Any]]:
        """Raccoglie articoli da feed RSS con filtering avanzato"""
        all_articles = []
        rss_feeds = self.config.get("ai_sources", {}).get("rss_feeds", [])
        max_per_source = self.config.get("selection_criteria", {}).get("max_articles_per_source", 5)
        
        logging.info(f"Raccogliendo articoli da {len(rss_feeds)} fonti RSS...")
        
        for feed_url in rss_feeds:
            try:
                logging.info(f"Fetching da: {feed_url}")
                feed = feedparser.parse(feed_url)
                
                # Estrai dominio per source tracking
                source_domain = urlparse(feed_url).netloc
                
                articles_from_source = 0
                for entry in feed.entries:
                    if articles_from_source >= max_per_source:
                        break
                    
                    # Crea oggetto articolo standardizzato
                    article = {
                        'title': entry.get('title', ''),
                        'link': entry.get('link', ''),
                        'summary': entry.get('summary', ''),
                        'published': entry.get('published', ''),
                        'source': source_domain,
                        'feed_url': feed_url,
                        'content': entry.get('content', [{}])[0].get('value', '') if entry.get('content') else ''
                    }
                    
                    # Verifica duplicati
                    article_hash = self.get_article_hash(article)
                    if article_hash in self.processed_articles:
                        continue
                    
                    # Applica filtri di qualità
                    if self.passes_quality_filters(article):
                        all_articles.append(article)
                        self.processed_articles.add(article_hash)
                        articles_from_source += 1
                        
                logging.info(f"Raccolti {articles_from_source} articoli da {source_domain}")
                
            except Exception as e:
                logging.error(f"Errore nel fetch da {feed_url}: {e}")
                continue
        
        logging.info(f"Totale articoli raccolti: {len(all_articles)}")
        return all_articles
    
    def passes_quality_filters(self, article: Dict[str, Any]) -> bool:
        """Applica filtri di qualità agli articoli"""
        criteria = self.config.get("selection_criteria", {})
        
        # Controlla lunghezza minima
        min_length = criteria.get("min_article_length", 200)
        content_length = len(article.get('summary', '') + article.get('content', ''))
        if content_length < min_length:
            return False
        
        # Controlla keywords AI
        ai_keywords = criteria.get("ai_keywords", [])
        exclude_keywords = criteria.get("exclude_keywords", [])
        
        text_content = f"{article.get('title', '')} {article.get('summary', '')}".lower()
        
        # Deve contenere almeno una keyword AI
        has_ai_keyword = any(keyword.lower() in text_content for keyword in ai_keywords)
        if not has_ai_keyword:
            return False
        
        # Non deve contenere keywords escluse
        has_excluded = any(keyword.lower() in text_content for keyword in exclude_keywords)
        if has_excluded:
            return False
        
        return True
    
    def fetch_arxiv_papers(self) -> List[Dict[str, Any]]:
        """Raccoglie paper da ArXiv (se abilitato)"""
        if not self.config.get("advanced_sources", {}).get("academic_papers", {}).get("enabled", False):
            return []
        
        papers = []
        categories = self.config.get("advanced_sources", {}).get("academic_papers", {}).get("arxiv_categories", [])
        max_papers = self.config.get("advanced_sources", {}).get("academic_papers", {}).get("max_papers_per_day", 2)
        
        logging.info(f"Raccogliendo paper da ArXiv: {categories}")
        
        for category in categories:
            try:
                # ArXiv API endpoint
                arxiv_url = f"http://export.arxiv.org/api/query?search_query=cat:{category}&start=0&max_results=5&sortBy=submittedDate&sortOrder=descending"
                
                response = requests.get(arxiv_url, timeout=30)
                response.raise_for_status()
                
                # Parse XML response (simplified)
                # In production, use xml.etree.ElementTree
                logging.info(f"Fetched ArXiv data for {category}")
                
            except Exception as e:
                logging.error(f"Errore fetch ArXiv {category}: {e}")
                continue
        
        return papers[:max_papers]
    
    def fetch_social_media_insights(self) -> List[Dict[str, Any]]:
        """Raccoglie insights da social media (se abilitato)"""
        if not self.config.get("advanced_sources", {}).get("social_media", {}).get("enabled", False):
            return []
        
        insights = []
        # Implementation for Twitter API, LinkedIn API, etc.
        logging.info("Social media insights non ancora implementati")
        return insights
    
    def fetch_news_api_articles(self) -> List[Dict[str, Any]]:
        """Raccoglie articoli da Google News API (se abilitato)"""
        if not self.config.get("advanced_sources", {}).get("api_sources", {}).get("google_news_api", False):
            return []
        
        articles = []
        # Implementation for Google News API
        logging.info("Google News API non ancora implementato")
        return articles
    
    def score_articles(self, articles: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Assegna punteggi agli articoli basati su criteri di rilevanza"""
        criteria = self.config.get("selection_criteria", {})
        freshness_weight = criteria.get("freshness_weight", 0.4)
        relevance_weight = criteria.get("relevance_weight", 0.6)
        ai_keywords = criteria.get("ai_keywords", [])
        
        for article in articles:
            # Score di freschezza (0-1)
            freshness_score = self.calculate_freshness_score(article.get('published', ''))
            
            # Score di rilevanza (0-1)
            relevance_score = self.calculate_relevance_score(article, ai_keywords)
            
            # Score finale
            final_score = (freshness_score * freshness_weight) + (relevance_score * relevance_weight)
            article['score'] = final_score
        
        # Ordina per punteggio decrescente
        return sorted(articles, key=lambda x: x.get('score', 0), reverse=True)
    
    def calculate_freshness_score(self, published_date: str) -> float:
        """Calcola punteggio di freschezza (articoli più recenti = punteggio maggiore)"""
        if not published_date:
            return 0.0
        
        try:
            # Parse della data (formato può variare)
            # Implementazione semplificata per evitare warning
            hours_ago = 24  # Assume 24 ore fa se non parseable
            
            # Score inversamente proporzionale alle ore passate
            max_hours = 168  # 7 giorni
            score = max(0, (max_hours - hours_ago) / max_hours)
            return score
            
        except Exception:
            return 0.5  # Default score
    
    def calculate_relevance_score(self, article: Dict[str, Any], ai_keywords: List[str]) -> float:
        """Calcola punteggio di rilevanza basato su keyword AI"""
        text_content = f"{article.get('title', '')} {article.get('summary', '')}".lower()
        
        # Conta keyword matches
        keyword_matches = sum(1 for keyword in ai_keywords if keyword.lower() in text_content)
        
        # Bonus per keyword nel titolo
        title_matches = sum(1 for keyword in ai_keywords if keyword.lower() in article.get('title', '').lower())
        
        # Score normalizzato
        total_keywords = len(ai_keywords)
        base_score = keyword_matches / total_keywords if total_keywords > 0 else 0
        title_bonus = title_matches * 0.2  # 20% bonus per keyword nel titolo
        
        return min(1.0, base_score + title_bonus)
    
    def select_top_articles(self, articles: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Seleziona i migliori articoli per la generazione post"""
        posts_per_day = self.config.get("schedule", {}).get("posts_per_day", 3)
        
        # Assegna punteggi
        scored_articles = self.score_articles(articles)
        
        # Seleziona i migliori
        selected = scored_articles[:posts_per_day]
        
        logging.info(f"Selezionati {len(selected)} articoli per generazione post:")
        for i, article in enumerate(selected, 1):
            logging.info(f"{i}. {article.get('title', 'N/A')} (Score: {article.get('score', 0):.2f})")
        
        return selected
    
    def collect_all_articles(self) -> List[Dict[str, Any]]:
        """Raccoglie articoli da tutte le fonti configurate"""
        all_articles = []
        
        # RSS Feeds (sempre attivi)
        rss_articles = self.fetch_rss_articles()
        all_articles.extend(rss_articles)
        
        # ArXiv Papers (se abilitato)
        arxiv_papers = self.fetch_arxiv_papers()
        all_articles.extend(arxiv_papers)
        
        # Social Media (se abilitato)
        social_insights = self.fetch_social_media_insights()
        all_articles.extend(social_insights)
        
        # News API (se abilitato)
        news_articles = self.fetch_news_api_articles()
        all_articles.extend(news_articles)
        
        logging.info(f"Totale articoli raccolti da tutte le fonti: {len(all_articles)}")
        return all_articles

# Esempio di utilizzo
if __name__ == "__main__":
    main()
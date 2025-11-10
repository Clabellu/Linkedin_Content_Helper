import feedparser
import os
import anthropic
from dotenv import load_dotenv
from newspaper import Article, ArticleException
# Importazioni per Tkinter
import tkinter as tk
from tkinter import simpledialog, messagebox # simpledialog non lo useremo subito, messagebox sì
import sys
from datetime import datetime 
import time
import re 
# import requests
import openai
import base64
import logging


from docx import Document

def setup_logging():
    """Setup del sistema di logging"""
    # Crea directory logs
    if not os.path.exists("logs"):
        os.makedirs("logs")
    
    # Nome file log con data
    date_str = datetime.now().strftime("%Y%m%d")
    log_file = f"logs/linkedin_generator_{date_str}.log"
    
    # Configura logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(funcName)s:%(lineno)d - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()  # Anche su console
        ]
    )


load_dotenv()
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")

if not ANTHROPIC_API_KEY:
    print("ATTENZIONE: La variabile d'ambiente ANTHROPIC_API_KEY non è impostata o non trovata nel file .env.")

if not OPENAI_API_KEY:
    print("ATTENZIONE: La variabile d'ambiente OPENAI_API_KEY non è impostata o non trovata nel file .env.")

FEED_FILE_PATH = "feeds.txt" # Definiamo il nome del file come costante

# Aggiungi questa costante vicino a FEED_FILE_PATH
PROCESSED_ARTICLES_FILE = "processed_articles.txt"

def load_processed_articles(filepath=PROCESSED_ARTICLES_FILE):
    """
    Carica gli URL degli articoli già processati e li restituisce come un 'set'
    per una ricerca veloce.
    """
    if not os.path.exists(filepath):
        return set() # Restituisce un set vuoto se il file non esiste
    
    with open(filepath, "r", encoding="utf-8") as f:
        # Crea un set di URL, rimuovendo spazi bianchi da ogni riga
        processed_urls = {line.strip() for line in f if line.strip()}
    print(f"Caricati {len(processed_urls)} URL di articoli già processati.")
    return processed_urls

def add_to_processed_articles(article_url, filepath=PROCESSED_ARTICLES_FILE):
    """
    Aggiunge un nuovo URL di un articolo processato al file di registro.
    Usa la modalità 'a' (append) per aggiungere in fondo al file senza cancellare.
    """
    try:
        with open(filepath, "a", encoding="utf-8") as f:
            f.write(article_url + "\n")
        print(f"Articolo {article_url} aggiunto al registro dei processati.")
        return True
    except Exception as e:
        print(f"Errore durante l'aggiornamento del registro dei processati: {e}")
        return False

def load_rss_feeds_from_file(filepath=FEED_FILE_PATH):
    """
    Carica la lista degli URL dei feed RSS da un file di testo.
    Ignora righe vuote e quelle che iniziano con '#'.
    """
    feeds = []
    if not os.path.exists(filepath):
        # Se il file non esiste, potremmo crearlo vuoto o restituire una lista vuota
        print(f"File dei feed '{filepath}' non trovato. Verrà creato se si aggiungono feed.")
        return [] # O potremmo crearlo: open(filepath, 'w').close()
    
    with open(filepath, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#"):
                feeds.append(line)
    print(f"Caricati {len(feeds)} feed da '{filepath}'")
    return feeds

def save_rss_feeds_to_file(feeds_list, filepath=FEED_FILE_PATH):
    """
    Salva la lista degli URL dei feed RSS su un file di testo, uno per riga.
    Sovrascrive il contenuto precedente del file.
    """
    try:
        with open(filepath, "w", encoding="utf-8") as f:
            for feed in feeds_list:
                f.write(feed + "\n")
        print(f"Salvati {len(feeds_list)} feed su '{filepath}'")
        return True
    except Exception as e:
        print(f"Errore durante il salvataggio dei feed su '{filepath}': {e}")
        messagebox.showerror("Errore Salvataggio Feed", f"Impossibile salvare i feed su file:\n{e}")
        return False
    
# All'avvio, carichiamo subito i feed. Questa variabile globale verrà usata e aggiornata.
# Rimuovi la vecchia definizione di rss_feeds = [...]
rss_feeds_list_global = load_rss_feeds_from_file()

def fetch_articles_from_feed(feed_url, limit=10):
    # ... (questa funzione rimane invariata) ...
    logging.info(f"Recupero articoli da: {feed_url} (limite: {limit})")
    try:
        feed = feedparser.parse(feed_url)
        articles = []
        entries_to_process = feed.entries[:limit]
        for entry in entries_to_process:
            article = {
                'title': entry.title,
                'link': entry.link,
                'published': entry.get('published', 'Data non disponibile'),
                'published_parsed': entry.get('published_parsed', None),
                'summary': entry.get('summary', 'Nessun riassunto disponibile')
            }
            articles.append(article)
        if not articles:
            logging.warning(f"Nessun articolo trovato per {feed_url}.")
        else:
            logging.info(f"Trovati {len(articles)} articoli in {feed_url}.")
        return articles
    except Exception as e:
        logging.error(f"Errore durante il recupero dal feed {feed_url}: {e}")
        return []


def sanitize_filename(filename):
    """
    Ripulisce una stringa per renderla un nome di file valido e più leggibile.
    Rimuove caratteri non alfanumerici (eccetto spazi, trattini, underscore),
    sostituisce gli spazi con underscore e limita la lunghezza.
    """
    filename = re.sub(r'[^\w\s-]', '', filename) # Rimuove caratteri non validi
    filename = re.sub(r'\s+', '_', filename).strip('_') # Sostituisce spazi con underscore
    return filename[:100] # Limita la lunghezza a 100 caratteri per sicurezza

def get_full_article_text_from_url(article_url):
    """
    Scarica e estrae il testo completo di un articolo da un URL usando newspaper3k.
    Restituisce il testo dell'articolo o None se l'estrazione fallisce.
    """
    if not article_url:
        return None
        
    logging.info(f"Estrazione testo da: {article_url}")
    try:
        # Impostiamo 'browser_user_agent' per sembrare un browser normale e ridurre errori
        article_parser = Article(article_url, language='it')
        article_parser.download()
        article_parser.parse()
        
        if article_parser.text:
            logging.info("Testo dell'articolo estratto con successo.")
            return article_parser.text
        else:
            logging.warning("Testo dell'articolo vuoto.")
            return None
    except ArticleException as e:
        logging.error(f"Errore durante il download o il parsing dell'articolo (newspaper3k): {e}")
        return None
    except Exception as e:
        logging.error(f"Errore generico durante l'estrazione del testo da {article_url}: {e}")
        return None
    
def generate_post_image_google(image_generation_prompt, output_filepath, api_key):
    """
    Genera un'immagine basata su un prompt testuale usando Google Gemini API
    """
    if not api_key:
        print("ERRORE: La chiave API di Google non è configurata.")
        return False

    print(f"\nRichiesta di generazione immagine a Google Gemini con prompt: '{image_generation_prompt}'")
    try:
        from google import genai
        from google.genai import types
        from PIL import Image
        from io import BytesIO
        
        # Configura il client Google Gemini
        genai.configure(api_key=api_key)
        client = genai.Client()
        
        # Genera l'immagine usando Gemini
        response = client.models.generate_content(
            model="gemini-2.5-flash-image-preview",
            contents=[image_generation_prompt],
        )
        
        # Gestione della risposta - estrai l'immagine generata
        if response and hasattr(response, 'candidates') and response.candidates:
            candidate = response.candidates[0]
            if hasattr(candidate, 'content') and candidate.content.parts:
                for part in candidate.content.parts:
                    if hasattr(part, 'inline_data') and part.inline_data:
                        # Decodifica i dati dell'immagine
                        image_data = part.inline_data.data
                        
                        # Se i dati sono in base64, decodificali
                        if isinstance(image_data, str):
                            image_bytes = base64.b64decode(image_data)
                        else:
                            image_bytes = image_data
                        
                        # Salva l'immagine
                        with open(output_filepath, 'wb') as f:
                            f.write(image_bytes)
                        
                        print(f"Immagine Google generata e salvata con successo in: {output_filepath}")
                        return True
        
        print("ERRORE: Risposta Google API non contiene immagine generata")
        return False
        
    except ImportError as e:
        print(f"ERRORE: Librerie Google Gemini non installate: {e}")
        print("Installare con: pip install google-generativeai")
        return False
    except Exception as e:
        print(f"ERRORE durante la generazione dell'immagine con Google: {e}")
        return False

def generate_post_image(image_generation_prompt, output_filepath, provider="openai"):
    """
    Genera un'immagine basata su un prompt testuale usando OpenAI o Google API
    """
    print(f"\nRichiesta di generazione immagine con {provider.upper()} con prompt: '{image_generation_prompt}'")
    
    if provider.lower() == "google":
        google_key = os.getenv('GOOGLE_API_KEY')
        if not google_key:
            print("ERRORE: La chiave API di Google non è configurata, provo con OpenAI...")
            provider = "openai"
        else:
            return generate_post_image_google(image_generation_prompt, output_filepath, google_key)
    
    if provider.lower() == "openai":
        if not OPENAI_API_KEY:
            print("ERRORE: La chiave API di OpenAI non è configurata.")
            return False
        
        try:
            client = openai.OpenAI(api_key=OPENAI_API_KEY)

            # Usa l'API corretta di OpenAI per la generazione di immagini
            response = client.images.generate(
                model="dall-e-3",
                prompt=image_generation_prompt,
                size="1024x1024",
                quality="standard",
                n=1
            )

            # Estrae l'URL dell'immagine generata
            image_url = response.data[0].url

            # Scarica l'immagine dall'URL
            import requests
            image_response = requests.get(image_url)
            image_response.raise_for_status()

            # Salva l'immagine nel file
            with open(output_filepath, 'wb') as f:
                f.write(image_response.content)

            print(f"Immagine generata e salvata con successo in: {output_filepath}")

            # Controlla se c'è un prompt rivisto
            if hasattr(response.data[0], 'revised_prompt') and response.data[0].revised_prompt:
                print(f"Prompt originale migliorato dall'IA: '{response.data[0].revised_prompt}'")

            return True

        except Exception as e:
            print(f"ERRORE durante la generazione dell'immagine con OpenAI: {e}")
            return False

    print(f"ERRORE: Provider non supportato: {provider}")
    return False

def generate_linkedin_post_with_claude(article_data, interactive_mode=True):

    article_title = article_data.get('title', 'Titolo non disponibile')
    article_link = article_data.get('link', '')

    logging.info(f"Inizio generazione post per: {article_title}")

    """
    Genera un post per LinkedIn E un prompt per l'immagine, poi chiama la funzione 
    per generare l'immagine e salva tutto.
    """
    if not ANTHROPIC_API_KEY:
        if interactive_mode: 
            messagebox.showerror("Errore API", "Chiave API di Anthropic non configurata.")
        return False
    
    # --- Recupero Contenuto ---
    logging.info(f"Recupero contenuto dell'articolo da: {article_link}")
    article_text_content = get_full_article_text_from_url(article_link) or article_data.get('summary', '')

    if not article_text_content:
        if interactive_mode: 
            messagebox.showerror("Errore Contenuto", f"Impossibile recuperare il contenuto per: {article_title}")
        return False

    # --- Prompt per Claude ---
    prompt_message = f"""Sei un esperto divulgatore e social media manager.
Basandoti sul seguente articolo intitolato "{article_title}", svolgi due compiti:

1. Scrivi un post per LinkedIn di 250-300 parole seguendo queste linee guida:

STRUTTURA DEL POST:
• Inizia con un HOOK coinvolgente (mai "Cari professionisti" o simili formalismi)
• Usa esempi di apertura come:
  - "Ieri ho scoperto qualcosa che ha cambiato il mio approccio a..."
  - "3 anni fa pensavo che [X] fosse impossibile. Mi sbagliavo."
  - "Ho appena testato [tecnologia/metodo] e i risultati mi hanno sorpreso"
  - "Quello che sto per condividere può farti risparmiare ore di lavoro"

TONO E STILE:
• Usa un tono conversazionale, come se stessi parlando con un collega
• Scrivi in prima persona quando possibile
• Includi 2-3 emoji pertinenti (non esagerare)
• Usa frasi breve e paragrafi di 1-2 righe per facilità di lettura

CONTENUTO:
• Racconta una storia o esperienza personale collegata al tema
• Spiega concetti complessi con analogie semplici
• Includi dati concreti o numeri quando disponibili
• Fornisci 2-3 takeaway pratici e actionable

ENGAGEMENT:
• Termina con una domanda aperta per stimolare i commenti
• Includi una call-to-action chiara
• Usa spazi bianchi per rendere il post più leggibile

• Aggiungi il link dell'articolo originale alla fine: {article_link}

2. Dopo il post, scrivi un prompt per immagine AI (in inglese, max 50 parole) che rappresenti visivamente il concetto principale del post.

Per favore, formatta la tua risposta esattamente così, senza aggiungere altro testo:
[POST]
(Qui il testo del post per LinkedIn)
[IMAGE_PROMPT]
(Qui il prompt per l'immagine in inglese)

--- TESTO ARTICOLO ---
{article_text_content}
---
"""
    
    # --- Chiamata a Claude e Gestione Risposta ---
    try:
        logging.info("Chiamata all'API di Claude per generare testo e prompt immagine...")
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        response = client.messages.create(
            model="claude-3-7-sonnet-20250219",
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt_message}]
        )
        
        if not (response.content and len(response.content) > 0 and hasattr(response.content[0], 'text')):
            logging.error("Risposta Claude vuota o formato non valido")
            if interactive_mode: 
                messagebox.showerror("Errore Risposta Claude", "La risposta da Claude è vuota o in un formato imprevisto.")
            return False

        # CORREZIONE: Definiamo la variabile con la risposta completa qui
        full_response_text = response.content[0].text
        
        # --- Estrazione del Post e del Prompt Immagine ---
        try:
            post_text = full_response_text.split("[IMAGE_PROMPT]")[0].replace("[POST]", "").strip()
            image_prompt_text = full_response_text.split("[IMAGE_PROMPT]")[1].strip()

            logging.info("Post generato con successo da Claude")
            logging.info(f"Prompt immagine: {image_prompt_text}")
        except IndexError:
            logging.error("Formato risposta Claude non valido - manacano i tag [POST] o [IMAGE_PROMPT]")
            if interactive_mode: 
                messagebox.showerror("Errore Formato Risposta", "Claude non ha restituito l'output nel formato [POST]...[IMAGE_PROMPT] atteso.")
            return False
            
        # --- Salvataggio File DOCX e Immagine PNG ---
        try:
            output_folder = "generated_posts"
            os.makedirs(output_folder, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            sanitized_title = sanitize_filename(article_title)
            
            # CORREZIONE: Definiamo i nomi dei file prima di usarli
            filename_base = f"{timestamp}_{sanitized_title}"
            filepath_docx = os.path.join(output_folder, f"{filename_base}.docx")
            filepath_image = os.path.join(output_folder, f"{filename_base}.png")

            # Salva il documento Word
            document = Document()
            document.add_heading('Post LinkedIn Generato', level=1)
            document.add_paragraph(post_text)
            # Potresti aggiungere altri dettagli come il link, ecc.
            document.save(filepath_docx) # CORREZIONE: Usa la variabile corretta
            logging.info(f"Post salvato con successo come file Word in: {filepath_docx}")

            # Genera e salva l'immagine
            logging.info("Generazione immagine in corso...")
            image_success = generate_post_image(image_prompt_text, filepath_image) # CORREZIONE: Usa la variabile corretta
            
            if image_success:
                logging.info(f"Immagine salvata con successo in: {filepath_image}")
            else:
                logging.warning("Generazione immagine fallita")

            if interactive_mode:
                messagebox.showinfo("Successo", "Post e immagine sono stati generati e salvati!")
            
            logging.info(f"Generazione completata con successo per: {article_title}")
            
            return True # Successo!
        
        except Exception as e_save:
            logging.error(f"Errore durante il salvataggio dei file: {e_save}")
            if interactive_mode: 
                messagebox.showerror("Errore Salvataggio", f"Errore durante il salvataggio dei file:\n{e_save}")
            return False

    except Exception as e:
        logging.error(f"Errore durante la chiamata a Claude: {e}")
        if interactive_mode: 
            messagebox.showerror("Errore Chiamata Claude", f"Errore durante la chiamata a Claude:\n{e}")
        return False
    
def run_automated_post_generation():
    """
    Esegue il ciclo automatico: calcola un 'punteggio di attualità' per ogni
    nuovo articolo e genera post per i 3 con il punteggio più alto.
    """
    logging.info("=== INIZIO ESECUZIONE AUTOMATICA ===")

    processed_urls = load_processed_articles()
    all_feeds = load_rss_feeds_from_file()

    if not all_feeds:
        logging.warning("Nessun feed RSS configurato.")
        return

    # 1. Raccogliamo tutti gli articoli e calcoliamo il loro punteggio di attualità
    all_new_articles_with_score = []
    now_timestamp = datetime.now().timestamp() # Timestamp attuale come riferimento

    print("\nRaccolta e valutazione di tutti i nuovi articoli...")
    for feed_url in all_feeds:
        articles_from_this_feed = fetch_articles_from_feed(feed_url, limit=10)
        
        # Usiamo enumerate per avere l'indice di ogni articolo nel suo feed (0, 1, 2...)
        for index, article in enumerate(articles_from_this_feed):
            article_link = article.get('link')
            
            # Se l'articolo ha un link e non è stato già processato...
            if article_link and article_link not in processed_urls:
                score = 0
                # Calcoliamo il punteggio
                if article.get('published_parsed'):
                    # Se ha una data, il punteggio è il suo timestamp
                    score = time.mktime(article['published_parsed'])
                else:
                    # Se non ha una data, il punteggio è basato sul tempo attuale
                    # meno una penalità per la sua posizione nel feed.
                    # Articoli più in alto (indice basso) avranno un punteggio più alto.
                    # La penalità (es. 3600 secondi = 1 ora) assicura che articoli
                    # dello stesso feed senza data siano ordinati correttamente.
                    penalty = index * 3600 
                    score = now_timestamp - penalty
                
                # Aggiungiamo alla lista l'articolo e il suo punteggio
                all_new_articles_with_score.append({'score': score, 'article': article})

    if not all_new_articles_with_score:
        print("\nNessun nuovo articolo trovato in nessun feed. Uscita.")
        return

    # 2. Ordiniamo la lista in base al punteggio, dal più alto al più basso
    logging.info(f"Trovati {len(all_new_articles_with_score)} nuovi articoli.")
    sorted_articles = sorted(
        all_new_articles_with_score,
        key=lambda x: x['score'],
        reverse=True
    )

    # 3. Prendiamo i primi 3 articoli da processare
    articles_to_process = sorted_articles[:1]
    
    print(f"Selezionati {len(articles_to_process)} articoli con punteggio più alto per la generazione dei post.")

    # 4. Eseguiamo un ciclo sugli articoli selezionati e generiamo i post
    for i, scored_item in enumerate(articles_to_process):
        article_to_process = scored_item['article'] # Estraiamo l'articolo vero e proprio
        print(f"\n--- Processo l'articolo #{i+1}/{len(articles_to_process)}: '{article_to_process['title']}' ---")
        
        success = generate_linkedin_post_with_claude(article_to_process, interactive_mode=False)

        if success:
            add_to_processed_articles(article_to_process['link'])
        else:
            logging.warning(f"Generazione o salvataggio del post per '{article_to_process['title']}' sono falliti.")

    logging.info("=== ESECUZIONE AUTOMATICA COMPLETATA ===")

def show_article_selection_window(initial_articles_list):
    global rss_feeds_list_global 

    window = tk.Tk()
    window.title("Selezione Articolo per Post LinkedIn")
    window.geometry("700x450") 

    manage_feeds_button = tk.Button(window, text="Gestisci Fonti RSS",
                                    command=lambda: open_manage_feeds_window(window),
                                    font=("Arial", 10), bg="lightyellow")
    manage_feeds_button.pack(pady=(10,0))

    label = tk.Label(window, text="Seleziona un articolo per generare un post:", font=("Arial", 14))
    label.pack(pady=5)

    listbox_articles = tk.Listbox(window, width=100, height=15, selectmode=tk.SINGLE, font=("Arial", 10))
    
    def on_generate_post_click():
        selected_indices = listbox_articles.curselection()
        if not selected_indices:
            messagebox.showwarning("Nessuna Selezione", "Per favore, seleziona un articolo dalla lista.", parent=window)
            return
        
        if not hasattr(window, 'current_display_articles') or not window.current_display_articles:
            messagebox.showerror("Errore", "Lista articoli non disponibile. Prova a ricaricare.", parent=window)
            return

        selected_index = selected_indices[0]
        if selected_index < len(window.current_display_articles):
            selected_article_data = window.current_display_articles[selected_index]
            print(f"\nArticolo selezionato per il post (dalla GUI): {selected_article_data['title']}")
            generate_linkedin_post_with_claude(selected_article_data, interactive_mode=True)
        else:
            messagebox.showerror("Errore Indice", "Errore nella selezione dell'articolo. Prova a ricaricare la lista.", parent=window)

    button_frame = tk.Frame(window)
    generate_button = tk.Button(button_frame, text="Genera Post da Articolo Selezionato", 
                                command=on_generate_post_click,
                                font=("Arial", 12), bg="lightgreen")
    generate_button.pack(side=tk.LEFT, padx=10)

    close_button = tk.Button(button_frame, text="Chiudi", command=window.destroy, font=("Arial", 12), bg="lightcoral")
    close_button.pack(side=tk.LEFT, padx=10)

    def populate_articles_listbox():
        listbox_articles.delete(0, tk.END)
        current_articles_to_display = []
        if not rss_feeds_list_global:
             messagebox.showinfo("Nessun Feed RSS", "Nessun feed RSS configurato. Aggiungine tramite 'Gestisci Fonti RSS'.", parent=window)
        else:
            print("\nAggiornamento articoli dai feed configurati...")
            for feed_url_item in rss_feeds_list_global:
                articles_from_current_feed = fetch_articles_from_feed(feed_url_item)
                current_articles_to_display.extend(articles_from_current_feed)
        
        if not current_articles_to_display:
            listbox_articles.insert(tk.END, "Nessun articolo trovato dai feed configurati.")
            generate_button.config(state=tk.DISABLED)
        else:
            for i, article in enumerate(current_articles_to_display):
                listbox_articles.insert(tk.END, f"{i+1}. {article['title']}")
            generate_button.config(state=tk.NORMAL)
        
        window.current_display_articles = current_articles_to_display

    reload_articles_button = tk.Button(window, text="Ricarica Articoli dai Feed",
                                     command=populate_articles_listbox,
                                     font=("Arial", 10))
    
    listbox_articles.pack(pady=10, padx=20, fill=tk.BOTH, expand=True)
    reload_articles_button.pack(pady=5)
    button_frame.pack(pady=10)

    populate_articles_listbox() 

    if not hasattr(window, 'current_display_articles') or not window.current_display_articles:
        generate_button.config(state=tk.DISABLED)
    else:
        generate_button.config(state=tk.NORMAL)

    window.mainloop() # <-- LA FUNZIONE DEVE TERMINARE QUI

# (Assicurati che 'import tkinter as tk' e 'from tkinter import messagebox, simpledialog' siano presenti)
# simpledialog ci servirà per l'input del nuovo URL

def open_manage_feeds_window(parent_window):
    global rss_feeds_list_global # Per poter modificare la lista globale dei feed

    manage_window = tk.Toplevel(parent_window)
    manage_window.title("Gestisci Fonti RSS")
    manage_window.geometry("600x400")
    manage_window.grab_set() # Rende questa finestra modale (blocca l'interazione con la finestra principale)

    # --- Elementi della Finestra di Gestione ---
    tk.Label(manage_window, text="URL dei Feed RSS Attuali:", font=("Arial", 12)).pack(pady=(10,0))

    listbox_feeds = tk.Listbox(manage_window, width=90, height=10, selectmode=tk.SINGLE)
    listbox_feeds.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

    def refresh_listbox():
        listbox_feeds.delete(0, tk.END) # Pulisce la lista
        for feed_url in rss_feeds_list_global:
            listbox_feeds.insert(tk.END, feed_url)

    refresh_listbox() # Popola la lista all'apertura

    # --- Funzioni per i Pulsanti ---
    def add_feed():
        new_url = simpledialog.askstring("Aggiungi Feed", "Inserisci il nuovo URL del Feed RSS:", parent=manage_window)
        if new_url and new_url.strip():
            new_url = new_url.strip()
            if new_url not in rss_feeds_list_global:
                rss_feeds_list_global.append(new_url)
                refresh_listbox()
                print(f"Feed aggiunto (non ancora salvato): {new_url}")
            else:
                messagebox.showwarning("Feed Esistente", "Questo URL è già presente nella lista.", parent=manage_window)
        elif new_url is not None: # L'utente ha premuto OK ma ha lasciato il campo vuoto
             messagebox.showwarning("URL Vuoto", "L'URL non può essere vuoto.", parent=manage_window)


    def remove_feed():
        selected_indices = listbox_feeds.curselection()
        if not selected_indices:
            messagebox.showwarning("Nessuna Selezione", "Seleziona un URL dalla lista per rimuoverlo.", parent=manage_window)
            return
        
        selected_index = selected_indices[0]
        url_to_remove = rss_feeds_list_global.pop(selected_index) # Rimuove dalla lista e ottiene l'elemento
        refresh_listbox()
        print(f"Feed rimosso (non ancora salvato): {url_to_remove}")

    def save_and_close():
        if save_rss_feeds_to_file(rss_feeds_list_global): # Salva la lista corrente su file
             messagebox.showinfo("Salvato", "La lista dei feed è stata salvata.", parent=manage_window)
        # Non chiudiamo subito, l'utente potrebbe voler fare altre modifiche
        # Per chiudere, l'utente può usare il pulsante "Chiudi Finestra" o la X
    
    def close_window_action():
        # Potremmo chiedere conferma se ci sono modifiche non salvate, per ora chiude e basta
        # La funzione main ricaricherà i feed dal file all'occorrenza
        manage_window.destroy()


    # --- Pulsanti ---
    buttons_frame = tk.Frame(manage_window)
    buttons_frame.pack(pady=10)

    tk.Button(buttons_frame, text="Aggiungi Feed", command=add_feed, bg="lightgreen").pack(side=tk.LEFT, padx=5)
    tk.Button(buttons_frame, text="Rimuovi Selezionato", command=remove_feed, bg="lightcoral").pack(side=tk.LEFT, padx=5)
    tk.Button(buttons_frame, text="Salva Modifiche su File", command=save_and_close, bg="lightblue").pack(side=tk.LEFT, padx=5)
    tk.Button(manage_window, text="Chiudi Finestra Gestione", command=close_window_action, font=("Arial", 10)).pack(pady=5)


def main():
    setup_logging()  # Setup del logging all'avvio
    logging.info("Avvio dell'applicazione LinkedIn Content Helper")
    show_article_selection_window(None)

if __name__ == "__main__":
    # Controlla se è stato passato l'argomento '--auto' dalla riga di comando
    if len(sys.argv) > 1 and sys.argv[1] == '--auto':
        run_automated_post_generation()
    else:
        # Se nessun argomento (o un argomento diverso) è presente, avvia la GUI
        main()